# -*- coding: utf-8 -*-
"""
把一份"内容规格(JSON)"渲染成排版固定、舒适不拥挤的中英逐句对照 Word 文档。

设计意图：
- 译文/题目等"内容"由 Claude 每次根据真题生成（写进 JSON），
- 而"排版"（字体、行距、留白、选项并排方式、正解红色✔等）由本脚本统一固定，
  这样每次产出的文档观感一致，不必反复重写格式代码。

用法：
    python render_bilingual_docx.py spec.json 输出.docx

spec.json 顶层字段：
    title    : 大标题（如 "2026 年（英语二）"）
    subtitle : 副标题（可选）
    note     : 顶部斜体小字说明（可选）
    blocks   : 内容块数组，元素 type 见下方 BLOCK TYPES

BLOCK TYPES（每个元素是一个 dict）：
  {"type":"h1","text":"Section I  Use of English  完形填空"}      一级标题
  {"type":"h2","text":"..."} / {"type":"h3","text":"..."}        二/三级标题
  {"type":"note","text":"..."}                                   斜体小字提示
  {"type":"pair","en":"英文句","zh":"中文句"}                     逐句对照（英一行、中一行，句对间留白）
  {"type":"cloze_opts","num":1,"opts":[["A","afraid","害怕的",false],
                                       ["B","critical","挑剔的",false],
                                       ["C","proud","自豪的",true],
                                       ["D","tolerant","宽容的",false]]}
        完形选项：4 个并排成一行（英一行、中一行），正确项 true → 红色加粗✔
  {"type":"qhead","num":21,"en":"题干英文","zh":"题干中文"}        阅读题干（加粗）
  {"type":"read_opts","opts":[[L,en,zh,correct],...]}            阅读选项：每行 2 个（英一行、中一行）
  {"type":"sub","letter":"A","en":"...","zh":"...","tag":"←【第42题 Christopher Kettle】✔"}
        Part B 待选小标题；tag 为标注（命中题用红色，无 tag 或写"（多余干扰项）"则为普通/斜体）
        若 "correct": true 则 tag 用红色加粗，否则普通。
  {"type":"numbered","num":41,"en":"英文条目","zh":"中文条目","ans":"F"}
        通用"编号条目 + 答案"。适配 Part B 的判断正误(T/F)、排序、句子填空等变体，
        以及翻译/写作里需要带编号的句子。ans 可省略；给出时红色加粗显示【答案：F】。
"""
import sys, json
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn

ZH_COLOR = RGBColor(0x1F, 0x4E, 0x79)   # 深蓝（中文）
ANS_COLOR = RGBColor(0xC0, 0x00, 0x00)  # 红（正确答案 / 命中标注）


def _set_cjk(run):
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def build(spec, out_path):
    doc = Document()

    # —— 适中页边距 ——
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Cm(2.2)
        sec.left_margin = sec.right_margin = Cm(2.4)

    # —— 默认字体 + 舒适行距 ——
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(11)
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.25

    # —— 标题区 ——
    doc.add_heading(spec['title'], level=0)
    if spec.get('subtitle'):
        doc.add_paragraph(spec['subtitle'])
    if spec.get('note'):
        p = doc.add_paragraph(); r = p.add_run(spec['note'])
        r.italic = True; r.font.size = Pt(9)

    def add_pair(en, zh):
        pe = doc.add_paragraph(); pe.add_run(en)
        pz = doc.add_paragraph(); pz.paragraph_format.space_after = Pt(9)
        rz = pz.add_run(zh); rz.font.color.rgb = ZH_COLOR; _set_cjk(rz)

    def add_cloze_opts(num, opts):
        pe = doc.add_paragraph()
        pe.add_run('%s. ' % num).bold = True
        for i, (L, e, z, c) in enumerate(opts):
            if i:
                pe.add_run('   ')
            r = pe.add_run('[%s] %s%s' % (L, e, '✔' if c else ''))
            if c:
                r.bold = True; r.font.color.rgb = ANS_COLOR
        pz = doc.add_paragraph(); pz.paragraph_format.space_after = Pt(9)
        for i, (L, e, z, c) in enumerate(opts):
            if i:
                pz.add_run('   ')
            r = pz.add_run('[%s] %s%s' % (L, z, '✔' if c else ''))
            _set_cjk(r)
            if c:
                r.bold = True; r.font.color.rgb = ANS_COLOR
            else:
                r.font.color.rgb = ZH_COLOR

    def add_qhead(num, en, zh):
        pe = doc.add_paragraph(); pe.paragraph_format.space_before = Pt(8)
        pe.add_run('%s. %s' % (num, en)).bold = True
        pz = doc.add_paragraph()
        rz = pz.add_run('     %s' % zh)
        rz.bold = True; rz.font.color.rgb = ZH_COLOR; _set_cjk(rz)

    def _row(p, items, use_zh):
        for i, (L, e, z, c) in enumerate(items):
            if i:
                p.add_run('    ')
            r = p.add_run('[%s] %s%s' % (L, z if use_zh else e, '✔' if c else ''))
            if use_zh:
                _set_cjk(r)
            if c:
                r.bold = True; r.font.color.rgb = ANS_COLOR
            elif use_zh:
                r.font.color.rgb = ZH_COLOR

    def add_read_opts(opts):
        for k in range(0, len(opts), 2):
            row = opts[k:k+2]
            _row(doc.add_paragraph(), row, False)
            pz = doc.add_paragraph(); pz.paragraph_format.space_after = Pt(4)
            _row(pz, row, True)

    def add_numbered(num, en, zh, ans):
        pe = doc.add_paragraph(); pe.paragraph_format.space_before = Pt(6)
        pe.add_run('%s. ' % num).bold = True
        pe.add_run(en)
        if ans:
            ra = pe.add_run('   【答案：%s】' % ans); ra.bold = True; ra.font.color.rgb = ANS_COLOR
        pz = doc.add_paragraph(); pz.paragraph_format.space_after = Pt(6)
        rz = pz.add_run('     %s' % zh); rz.font.color.rgb = ZH_COLOR; _set_cjk(rz)

    def add_sub(letter, en, zh, tag, correct):
        pe = doc.add_paragraph(); pe.paragraph_format.space_before = Pt(4)
        pe.add_run('[%s] %s' % (letter, en))
        if tag:
            rt = pe.add_run('   ' + tag)
            if correct:
                rt.bold = True; rt.font.color.rgb = ANS_COLOR
            else:
                rt.italic = True
        pz = doc.add_paragraph()
        rz = pz.add_run('     [%s] %s' % (letter, zh)); rz.font.color.rgb = ZH_COLOR; _set_cjk(rz)

    for b in spec['blocks']:
        t = b['type']
        if t == 'h1':
            doc.add_heading(b['text'], level=1)
        elif t == 'h2':
            doc.add_heading(b['text'], level=2)
        elif t == 'h3':
            doc.add_heading(b['text'], level=3)
        elif t == 'note':
            p = doc.add_paragraph(); r = p.add_run(b['text']); r.italic = True; r.font.size = Pt(9)
        elif t == 'pair':
            add_pair(b['en'], b['zh'])
        elif t == 'cloze_opts':
            add_cloze_opts(b['num'], b['opts'])
        elif t == 'qhead':
            add_qhead(b['num'], b['en'], b['zh'])
        elif t == 'read_opts':
            add_read_opts(b['opts'])
        elif t == 'sub':
            add_sub(b['letter'], b['en'], b['zh'], b.get('tag', ''), b.get('correct', False))
        elif t == 'numbered':
            add_numbered(b['num'], b['en'], b['zh'], b.get('ans', ''))
        else:
            raise ValueError('unknown block type: %s' % t)

    doc.save(out_path)
    print('saved:', out_path)


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print('usage: python render_bilingual_docx.py spec.json out.docx'); sys.exit(1)
    with open(sys.argv[1], encoding='utf-8') as f:
        spec = json.load(f)
    build(spec, sys.argv[2])
